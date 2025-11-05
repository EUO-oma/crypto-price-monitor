import tkinter as tk
from tkinter import filedialog
from docx import Document

# Create the main window
win = tk.Tk()
win.title("Word Reader")
win.geometry("800x600")

# Create a menu bar with "File", "View", and "Settings" menus
menu_bar = tk.Menu(win)

# Define a function to handle file opening
def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
    if file_path:
        text.delete("1.0", tk.END)
        document = Document(file_path)
        for para in document.paragraphs:
            text.insert(tk.END, para.text + "\n")

        # Save the file path as the last opened file
        with open("last_opened.txt", "w") as file:
            file.write(file_path)

# Define a function to change the font size
def change_font_size(size):
    text.config(font=("Helvetica", size))

# Define a function to switch to dark mode
def enable_dark_mode():
    win.config(bg="black")
    text.config(bg="black", fg="white")

# Define a function to switch to light mode
def disable_dark_mode():
    win.config(bg="white")
    text.config(bg="white", fg="black")

# Define a function to change the column spacing
def change_column_spacing(spacing):
    text.config(tabs=(spacing,))

    # Save the column spacing value
    with open("settings.txt", "w") as file:
        file.write(str(spacing))

# Create a text widget to display the contents of the Word document
text = tk.Text(win)
text.pack(fill=tk.BOTH, expand=True)


# Load the column spacing value from the settings file
try:
    with open("settings.txt", "r") as file:
        column_spacing_str = file.readline().strip()
        if column_spacing_str:
            column_spacing = int(column_spacing_str)
            text.tag_configure("col", lmargin1=0, lmargin2=column_spacing)
except (FileNotFoundError, ValueError):
    pass


# Create a "File" menu
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Open", command=open_file)
file_menu.add_command(label="Exit", command=win.quit)
menu_bar.add_cascade(label="File", menu=file_menu)

# Create a "View" menu
view_menu = tk.Menu(menu_bar, tearoff=0)
view_menu.add_command(label="Dark Mode", command=enable_dark_mode)
view_menu.add_command(label="Light Mode", command=disable_dark_mode)
view_menu.add_separator()

# Create a "Column Spacing" submenu
column_spacing_menu = tk.Menu(view_menu, tearoff=0)
for spacing in [0, 10, 20, 30, 40]:
    column_spacing_menu.add_command(label=str(spacing), command=lambda spacing=spacing: change_column_spacing(spacing))
view_menu.add_cascade(label="Column Spacing", menu=column_spacing_menu)

menu_bar.add_cascade(label="View", menu=view_menu)


# Create a "Settings" menu
settings_menu = tk.Menu(menu_bar, tearoff=0)

# Create a "Font Size" submenu
font_size_menu = tk.Menu(settings_menu, tearoff=0)
for font_size in [12, 14, 16, 18, 20, 100, 200]:
    font_size_menu.add_command(label=str(font_size), command=lambda size=font_size: change_font_size(size))
settings_menu.add_cascade(label="Font Size", menu=font_size_menu)

# Create a "Column Spacing" submenu
column_spacing_menu = tk.Menu(settings_menu, tearoff=0)
for spacing in [0, 10, 20, 30, 40]:
    column_spacing_menu.add_command(label=str(spacing), command=lambda spacing=spacing: change_column_spacing(spacing))
settings_menu.add_cascade(label="Column Spacing", menu=column_spacing_menu)

# Create a "Dark Mode" submenu
dark_mode_menu = tk.Menu(settings_menu, tearoff=0)
dark_mode_menu.add_command(label="Enable", command=enable_dark_mode)
dark_mode_menu.add_command(label="Disable", command=disable_dark_mode)
settings_menu.add_cascade(label="Dark Mode", menu=dark_mode_menu)

# Add the "Settings" menu to the menu bar
menu_bar.add_cascade(label="Settings", menu=settings_menu)


# Define a function to change the column spacing
def change_column_spacing(spacing):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = None  # Remove the existing font size
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')  # Set the font
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Inches(spacing/72)  # Convert from points to inches
    document.save("last_opened.docx")  # Save the changes to the last opened file

    # Save the column spacing value
    with open("settings.txt", "w") as file:
        file.write(str(spacing))

# Load the last opened file, if any
try:
    with open("last_opened.txt", "r") as file:
        last_file_path = file.readline().strip()  # Read the first line, which is the file path
        column_spacing = int(file.readline())  # Read the second line, which is the column spacing value
        if last_file_path:
            document = Document(last_file_path)
            change_column_spacing(column_spacing)  # Apply the column spacing
except FileNotFoundError:
    pass

# Create a "File" menu
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Open", command=open_file)
file_menu.add_command(label="Exit", command=win.quit)
menu_bar.add_cascade(label="File", menu=file_menu)

# Create a "View" menu
view_menu = tk.Menu(menu_bar, tearoff=0)
view_menu.add_command(label="Dark Mode", command=enable_dark_mode)
view_menu.add_command(label="Light Mode", command=disable_dark_mode)
view_menu.add_separator()

# Create a "Settings" menu
settings_menu = tk.Menu(menu_bar, tearoff=0)

# Create a "Font Size" submenu
font_size_menu = tk.Menu(settings_menu, tearoff=0)
for font_size in [12, 14, 16, 18, 20, 100, 200]:
    font_size_menu.add_command(label=str(font_size), command=lambda size=font_size: change_font_size(size))
settings_menu.add_cascade(label="Font Size", menu=font_size_menu)

# Create a "Column Spacing" submenu
column_spacing_menu = tk.Menu(settings_menu, tearoff=0)
for spacing in [0, 10, 20, 30, 40]:
    column_spacing_menu.add_command(label=str(spacing), command=lambda spacing=spacing: change_column_spacing(spacing))
settings_menu.add_cascade(label="Column Spacing", menu=column_spacing_menu)

# Add the "Settings" menu to the menu bar
menu_bar.add_cascade(label="Settings", menu=settings_menu)

# Add the menu bar to the window
win.config(menu=menu_bar)

# Start the Tkinter event loop
win.mainloop()
